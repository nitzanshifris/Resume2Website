"""
Text Extractor Service for CV2WEB MVP
Smart extraction with local processing first, then Google Vision OCR, fallback to AWS Textract
"""
import logging
from pathlib import Path
from typing import Optional, Tuple
import os
import unicodedata
import re
from services.local.keychain_manager import get_google_credentials_path, get_aws_credentials

# Document processing libraries
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
from striprtf.striprtf import rtf_to_text

# OCR libraries
try:
    from google.cloud import vision
    GOOGLE_VISION_AVAILABLE = True
except ImportError:
    GOOGLE_VISION_AVAILABLE = False
    
try:
    import boto3
    AWS_TEXTRACT_AVAILABLE = True
except ImportError:
    AWS_TEXTRACT_AVAILABLE = False

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Smart text extractor with three-tier approach:
    1. Local extraction (free & fast) 
    2. Google Cloud Vision OCR (primary OCR)
    3. AWS Textract (fallback OCR)
    """
    
    def __init__(self):
        # Initialize Google Vision client
        self.vision_client = None
        if GOOGLE_VISION_AVAILABLE:
            try:
                # Get credentials path from Keychain
                creds_path = get_google_credentials_path()
                if creds_path and os.path.exists(creds_path):
                    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = creds_path
                    self.vision_client = vision.ImageAnnotatorClient()
                    logger.info("Google Cloud Vision client initialized successfully")
                else:
                    logger.warning("Google credentials not found in Keychain")
            except Exception as e:
                logger.warning(f"Failed to initialize Google Vision: {e}")
                
        # Initialize AWS Textract client  
        self.textract_client = None
        if AWS_TEXTRACT_AVAILABLE:
            try:
                # Get AWS credentials from Keychain
                aws_creds = get_aws_credentials()
                if aws_creds.get('aws_access_key_id') and aws_creds.get('aws_secret_access_key'):
                    self.textract_client = boto3.client(
                        "textract",
                        **aws_creds
                    )
                    logger.info("AWS Textract client initialized successfully")
                else:
                    logger.warning("AWS credentials not found in Keychain")
            except Exception as e:
                logger.warning(f"Failed to initialize AWS Textract: {e}")
    
    def _normalize_text(self, text: str) -> str:
        """
        Normalize Unicode text to fix common issues:
        - Replace ligatures (fi, fl, etc.) with regular characters
        - Replace curly quotes with straight quotes
        - Normalize spaces and special characters
        - Remove zero-width characters
        """
        if not text:
            return text
            
        # First, handle ligatures explicitly before normalization
        ligature_replacements = {
            '\ufb00': 'ff',  # ff ligature
            '\ufb01': 'fi',  # fi ligature
            '\ufb02': 'fl',  # fl ligature
            '\ufb03': 'ffi', # ffi ligature
            '\ufb04': 'ffl', # ffl ligature
            '\ufb05': 'st',  # st ligature
            '\ufb06': 'st',  # st ligature (alternative)
        }
        
        for ligature, replacement in ligature_replacements.items():
            text = text.replace(ligature, replacement)
            
        # Now normalize using Unicode NFKC (compatibility composition)
        # This handles other Unicode normalization
        text = unicodedata.normalize('NFKC', text)
        
        # Replace common problematic characters
        replacements = {
            '\u2019': "'",  # Right single quotation mark → apostrophe
            '\u2018': "'",  # Left single quotation mark → apostrophe
            '\u201C': '"',  # Left double quotation mark
            '\u201D': '"',  # Right double quotation mark
            '\u2013': '-',  # En dash
            '\u2014': '-',  # Em dash
            '\u2026': '...',  # Horizontal ellipsis
            '\u00A0': ' ',  # Non-breaking space
            '\u200B': '',  # Zero-width space
            '\u200C': '',  # Zero-width non-joiner
            '\u200D': '',  # Zero-width joiner
            '\uFEFF': '',  # Zero-width no-break space (BOM)
        }
        
        for old_char, new_char in replacements.items():
            text = text.replace(old_char, new_char)
        
        # Fix spacing around percentages (e.g., "128 %" → "128%")
        text = re.sub(r'(\d+)\s+%', r'\1%', text)
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Trim each line
        lines = text.split('\n')
        text = '\n'.join(line.strip() for line in lines if line.strip())
        
        return text
    
    def extract_text(self, file_path: str) -> str:
        """
        Main extraction method with smart routing
        
        Args:
            file_path: Path to the file to extract text from
            
        Returns:
            Extracted text as string (never None)
        """
        path = Path(file_path)
        
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return ""
            
        # Get file extension
        file_ext = path.suffix.lower()
        logger.info(f"Extracting text from {path.name} (type: {file_ext})")
        
        # Step 1: Try local extraction first (free & fast)
        text, needs_ocr = self._try_local_extraction(path, file_ext)
        
        if text and not needs_ocr:
            logger.info(f"Local extraction successful: {len(text)} characters")
            return self._normalize_text(text)
            
        # Step 2: Use OCR if needed (images or failed extraction)
        if needs_ocr or not text:
            logger.info("Local extraction insufficient, attempting OCR...")
            text = self._extract_with_ocr(path)
            
        return self._normalize_text(text) if text else ""
    
    def _try_local_extraction(self, path: Path, file_ext: str) -> Tuple[str, bool]:
        """
        Try to extract text locally based on file type
        
        Returns:
            (extracted_text, needs_ocr) tuple
        """
        # Image files always need OCR
        image_extensions = {'.jpg', '.jpeg', '.png', '.webp', '.heic', '.heif', '.tiff', '.tif', '.bmp'}
        if file_ext in image_extensions:
            return "", True
            
        # Try document extraction
        try:
            if file_ext == '.pdf':
                return self._extract_pdf(path), False
            elif file_ext in ['.docx', '.doc']:
                return self._extract_docx(path), False
            elif file_ext == '.txt':
                return self._extract_txt(path), False
            elif file_ext == '.md':
                return self._extract_markdown(path), False
            elif file_ext == '.rtf':
                return self._extract_rtf(path), False
            elif file_ext in ['.html', '.htm']:
                return self._extract_html(path), False
            elif file_ext == '.odt':
                logger.warning("ODT files not supported in MVP, will use OCR")
                return "", True
            else:
                logger.warning(f"Unknown file type {file_ext}, will use OCR")
                return "", True
                
        except Exception as e:
            logger.error(f"Local extraction failed: {e}")
            return "", True
    
    def _extract_with_ocr(self, path: Path) -> str:
        """
        Extract text using OCR with fallback chain:
        Google Vision -> AWS Textract -> Empty string
        """
        # Try Google Vision first
        if self.vision_client:
            try:
                text = self._extract_with_google_vision(path)
                if text:
                    logger.info(f"Google Vision OCR successful: {len(text)} characters")
                    return text
            except Exception as e:
                logger.error(f"Google Vision failed: {e}")
                
        # Fallback to AWS Textract
        if self.textract_client:
            try:
                text = self._extract_with_textract(path)
                if text:
                    logger.info(f"AWS Textract OCR successful: {len(text)} characters")
                    return text
            except Exception as e:
                logger.error(f"AWS Textract failed: {e}")
                
        logger.error("All OCR methods failed")
        return ""
    
    # === Local Extraction Methods ===
    
    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF using PyPDF2"""
        text_parts = []
        
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num + 1}: {e}")
                    
        text = "\n\n".join(text_parts)
        
        # If PDF has no extractable text, it's probably scanned
        if not text.strip():
            logger.info("PDF appears to be scanned (no extractable text)")
            raise Exception("PDF needs OCR")
            
        return text
    
    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX file"""
        doc = Document(path)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
                
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        return "\n\n".join(text_parts)
    
    def _extract_txt(self, path: Path) -> str:
        """Extract text from plain text file"""
        # Try multiple encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise Exception("Failed to decode text file")
    
    def _extract_markdown(self, path: Path) -> str:
        """Extract text from Markdown (simple approach for MVP)"""
        import re
        text = path.read_text(encoding='utf-8')
        
        # Remove common markdown syntax
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)  # Headers
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Italic
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Links
        text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)  # Code blocks
        text = re.sub(r'`([^`]+)`', r'\1', text)  # Inline code
        
        return text
    
    def _extract_rtf(self, path: Path) -> str:
        """Extract text from RTF file"""
        rtf_content = path.read_text(encoding='utf-8', errors='ignore')
        return rtf_to_text(rtf_content)
    
    def _extract_html(self, path: Path) -> str:
        """Extract text from HTML file"""
        html_content = path.read_text(encoding='utf-8', errors='ignore')
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for element in soup(['script', 'style']):
            element.decompose()
            
        # Get text
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    
    # === OCR Methods ===
    
    def _extract_with_ocr(self, path: Path) -> str:
        """
        Extract text using OCR with fallback chain:
        Google Vision -> AWS Textract -> Empty string
        """
        # Check file size before OCR
        file_size_mb = path.stat().st_size / (1024 * 1024)
        if file_size_mb > 20:  # Google Vision limit
            logger.error(f"File too large for OCR: {file_size_mb:.1f}MB (max 20MB)")
            return ""
        
        # Try Google Vision first
        if self.vision_client:
            try:
                text = self._extract_with_google_vision(path)
                if text:
                    logger.info(f"Google Vision OCR successful: {len(text)} characters")
                    return text
            except Exception as e:
                logger.error(f"Google Vision failed: {e}")
                
        # Fallback to AWS Textract (max 5MB for sync API)
        if self.textract_client and file_size_mb <= 5:
            try:
                text = self._extract_with_textract(path)
                if text:
                    logger.info(f"AWS Textract OCR successful: {len(text)} characters")
                    return text
            except Exception as e:
                logger.error(f"AWS Textract failed: {e}")
        elif file_size_mb > 5:
            logger.warning("File too large for AWS Textract sync API (>5MB)")
                
        logger.error("All OCR methods failed")
        return ""
    
    def _extract_with_google_vision(self, path: Path) -> str:
        """Extract text using Google Cloud Vision API"""
        with open(path, 'rb') as image_file:
            content = image_file.read()
            
        image = vision.Image(content=content)
        response = self.vision_client.text_detection(image=image)
        
        if response.error.message:
            raise Exception(f"Google Vision API error: {response.error.message}")
            
        # Get full text from the first annotation
        if response.text_annotations:
            return response.text_annotations[0].description
            
        return ""
    
    def _extract_with_textract(self, path: Path) -> str:
        """Extract text using AWS Textract"""
        with open(path, 'rb') as document_file:
            document_bytes = document_file.read()
            
        # For PDFs larger than 5MB, would need to use async operation
        # For MVP, we'll use sync operation
        response = self.textract_client.detect_document_text(
            Document={'Bytes': document_bytes}
        )
        
        # Extract text from response
        text_parts = []
        for item in response.get("Blocks", []):
            if item["BlockType"] == "LINE":
                text_parts.append(item["Text"])
                
        return "\n".join(text_parts)


# === Create singleton instance ===
text_extractor = TextExtractor()


# === Convenience function for simple usage ===
def extract_text(file_path: str) -> str:
    """
    Simple function interface for text extraction
    
    Usage:
        text = extract_text("path/to/file.pdf")
    """
    return text_extractor.extract_text(file_path)


# === For testing ===
if __name__ == "__main__":
    import sys
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        print(f"Extracting text from: {test_file}")
        
        text = extract_text(test_file)
        print(f"\nExtracted {len(text)} characters")
        print(f"First 500 chars:\n{text[:500]}...")